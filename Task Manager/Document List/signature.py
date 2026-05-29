# =========================================================
# Import Required Libraries
# =========================================================

# Used to create and edit Word (.docx) documents
from docx import Document

# Used to set image size inside the Word document
from docx.shared import Inches

# Optional:
# Used to work with images (not required here but useful)
from PIL import Image


# =========================================================
# File Paths
# =========================================================

# Path of the image you want to insert
img_path = "image.png"

# Name of the Word document to create
doc_path = "converted_image_document.docx"


# =========================================================
# Create a New Word Document
# =========================================================

# Create an empty Word document object
doc = Document()


# =========================================================
# Add Title Heading
# =========================================================

# Add a heading to the document
doc.add_heading("Converted Image Document", level=1)


# =========================================================
# Extracted Text Content
# =========================================================

# Multi-line text content
text = """
7. Certification of the supplement:

7.1 Date: 01.07.2023

7.2 Given name(s), family name(s), signature(s):
Deniss Djakons

7.3 Position(s) of the person(s), certifying the Supplement:
Rector

7.4 Official stamp or seal:
"""


# =========================================================
# Add Text into Word Document
# =========================================================

# Insert paragraph text into the document
doc.add_paragraph(text)


# =========================================================
# Add Image into Word Document
# =========================================================

# Insert image with custom width
doc.add_picture(img_path, width=Inches(6.5))


# =========================================================
# Save Word Document
# =========================================================

# Save the document to the specified path
doc.save(doc_path)


# =========================================================
# Success Message
# =========================================================

print("Word document created successfully!")
print(f"Saved at: {doc_path}")